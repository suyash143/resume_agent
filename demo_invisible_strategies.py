#!/usr/bin/env python3.10
"""
Demo script showing the invisible keyword embedding strategies
"""

import docx
from docx.oxml import OxmlElement
from docx.shared import RGBColor, Pt

def create_demo_document():
    """Create a demo document showing invisible strategies"""
    print("🎯 Creating Demo Document with Invisible Strategies")
    print("=" * 50)
    
    # Create a new document
    doc = docx.Document()
    
    # Add title
    title = doc.add_heading('Resume Optimization Demo', 0)
    
    # Add visible content
    doc.add_paragraph("This is a sample resume with normal, visible content.")
    doc.add_paragraph("Skills: Communication, Leadership, Problem Solving")
    
    # Strategy 1: Document Metadata
    print("📝 Strategy 1: Adding keywords to document metadata...")
    keywords = ["python", "aws", "machine learning", "docker", "kubernetes", "ai", "nlp"]
    doc.core_properties.keywords = ", ".join(keywords[:10])
    doc.core_properties.comments = "ATS Keywords: " + ", ".join(keywords[10:])
    doc.core_properties.subject = "Software Engineer Resume"
    
    # Strategy 2: Hidden Text (completely invisible)
    print("👻 Strategy 2: Adding completely hidden text...")
    hidden_para = doc.add_paragraph()
    hidden_run = hidden_para.add_run(" ".join(keywords))
    r = hidden_run._r
    rPr = r.get_or_add_rPr()
    vanish = OxmlElement('w:vanish')
    rPr.append(vanish)
    
    # Strategy 3: White text on white background
    print("⚪ Strategy 3: Adding white text (invisible but searchable)...")
    white_para = doc.add_paragraph()
    white_run = white_para.add_run("Additional keywords: tensorflow pytorch fastapi langchain")
    white_run.font.color.rgb = RGBColor(255, 255, 255)  # White text
    white_run.font.size = Pt(1)  # Tiny font
    
    # Strategy 4: Custom document properties
    print("🔧 Strategy 4: Adding custom document properties...")
    try:
        custom_props = doc.custom_properties
        custom_props.add('ats_optimization', 'enabled')
        custom_props.add('target_keywords', ', '.join(keywords))
    except:
        print("   ⚠️  Custom properties not supported in this version")
    
    # Add more visible content
    doc.add_paragraph("Experience: 5+ years in software development")
    doc.add_paragraph("Education: Computer Science Degree")
    
    # Save the demo document
    output_path = "demo_invisible_strategies.docx"
    doc.save(output_path)
    
    print(f"\n✅ Demo document created: {output_path}")
    print(f"\n🔍 What happened:")
    print(f"   • Document looks normal to human readers")
    print(f"   • {len(keywords)} keywords embedded invisibly")
    print(f"   • ATS systems will detect all keywords")
    print(f"   • Metadata contains additional keywords")
    print(f"   • Hidden text is completely invisible")
    print(f"   • White text blends with background")
    
    return output_path

def analyze_demo_document(doc_path):
    """Analyze the demo document to show what was embedded"""
    print(f"\n🔬 Analyzing Demo Document: {doc_path}")
    print("=" * 40)
    
    try:
        doc = docx.Document(doc_path)
        
        # Check metadata
        print(f"📋 Document Metadata:")
        print(f"   Keywords: {doc.core_properties.keywords}")
        print(f"   Comments: {doc.core_properties.comments}")
        print(f"   Subject: {doc.core_properties.subject}")
        
        # Count paragraphs
        visible_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                visible_text.append(para.text.strip())
        
        print(f"\n📄 Visible Content ({len(visible_text)} paragraphs):")
        for i, text in enumerate(visible_text, 1):
            print(f"   {i}. {text}")
        
        print(f"\n💡 Key Points:")
        print(f"   • Human readers see only the visible content")
        print(f"   • ATS systems also scan metadata and hidden elements")
        print(f"   • Keywords are strategically distributed")
        print(f"   • No visual layout changes")
        
    except Exception as e:
        print(f"❌ Error analyzing document: {e}")

def main():
    """Run the demo"""
    try:
        # Create demo document
        demo_path = create_demo_document()
        
        # Analyze the document
        analyze_demo_document(demo_path)
        
        print(f"\n🎉 Demo Complete!")
        print(f"\n📖 To see this in action:")
        print(f"   1. Open {demo_path} in Microsoft Word")
        print(f"   2. The document looks normal")
        print(f"   3. But ATS systems will find all embedded keywords")
        print(f"   4. Check File > Properties to see metadata")
        
    except Exception as e:
        print(f"❌ Demo error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
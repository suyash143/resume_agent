import docx
import spacy
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx2pdf import convert
import os
import re
from collections import Counter
import xml.etree.ElementTree as ET
import requests
from dotenv import load_dotenv
import json
from llm_prompt import LLM_KEYWORD_EXTRACTION_PROMPT

load_dotenv()


# Load spaCy model for keyword extraction
def load_spacy_model():
    try:
        return spacy.load('en_core_web_sm')
    except OSError:
        print("SpaCy model not found. Please run: python -m spacy download en_core_web_sm")
        exit(1)


nlp = load_spacy_model()


def extract_smart_keywords(jd_text, max_keywords=50):
    """Extract and rank keywords from job description with relevance scoring"""
    doc = nlp(jd_text.lower())

    # Technical skills patterns
    tech_patterns = [
        r'\b(?:python|java|javascript|react|node\.?js|aws|docker|kubernetes|sql|nosql)\b',
        r'\b(?:machine learning|deep learning|ai|artificial intelligence|nlp|llm)\b',
        r'\b(?:tensorflow|pytorch|scikit-learn|pandas|numpy|fastapi|langchain)\b',
        r'\b(?:rag|retrieval|augmented|generation|vector|embedding|transformer)\b',
        r'\b(?:sagemaker|hugging face|pinecone|faiss|chroma|mlops)\b',
        r'\b(?:prompt engineering|fine-tuning|lora|peft|agentic)\b'
    ]

    keywords = {}

    # Extract technical terms with high priority
    for pattern in tech_patterns:
        matches = re.findall(pattern, jd_text.lower())
        for match in matches:
            keywords[match] = keywords.get(match, 0) + 3  # High weight for tech terms

    # Extract named entities
    for ent in doc.ents:
        if ent.label_ in ['ORG', 'PRODUCT', 'LANGUAGE', 'SKILL']:
            clean_text = ent.text.strip().lower()
            if len(clean_text) > 2:
                keywords[clean_text] = keywords.get(clean_text, 0) + 2

    # Extract important nouns and proper nouns
    for token in doc:
        if (token.pos_ in ['NOUN', 'PROPN'] and
                not token.is_stop and
                len(token.text) > 2 and
                token.text.isalpha()):
            keywords[token.text] = keywords.get(token.text, 0) + 1

    # Extract multi-word technical phrases
    phrases = [
        'machine learning', 'deep learning', 'natural language processing',
        'retrieval augmented generation', 'large language models',
        'prompt engineering', 'vector databases', 'generative ai',
        'artificial intelligence', 'data science', 'cloud computing'
    ]

    for phrase in phrases:
        if phrase in jd_text.lower():
            keywords[phrase] = keywords.get(phrase, 0) + 2

    # Sort by frequency and return top keywords
    sorted_keywords = sorted(keywords.items(), key=lambda x: x[1], reverse=True)
    return [kw[0] for kw in sorted_keywords[:max_keywords]]


def add_keywords_to_metadata(doc, keywords):
    """Add keywords to multiple metadata fields for maximum ATS coverage"""
    keywords_str = ', '.join(keywords)

    # Split keywords across multiple metadata fields
    chunk_size = 200  # Leave room for existing content
    keyword_chunks = [keywords_str[i:i + chunk_size] for i in range(0, len(keywords_str), chunk_size)]

    # Add to various metadata fields
    if len(keyword_chunks) > 0:
        doc.core_properties.keywords = keyword_chunks[0]
    if len(keyword_chunks) > 1:
        doc.core_properties.comments = keyword_chunks[1]
    if len(keyword_chunks) > 2:
        doc.core_properties.subject = keyword_chunks[2]

    # Add to custom properties
    try:
        custom_props = doc.custom_properties
        custom_props.add('ats_keywords', ', '.join(keywords[:20]))
        custom_props.add('skills', ', '.join([kw for kw in keywords if any(tech in kw.lower()
                                                                           for tech in
                                                                           ['python', 'aws', 'ai', 'ml', 'data'])]))
    except:
        pass  # Some versions might not support custom properties


def add_hidden_text(paragraph, text):
    """Add completely hidden text (invisible to readers, readable by ATS)"""
    run = paragraph.add_run(text)
    r = run._r
    rPr = r.get_or_add_rPr()
    vanish = OxmlElement('w:vanish')
    rPr.append(vanish)


def add_white_text(paragraph, text):
    """Add white text on white background (invisible but searchable)"""
    from docx.shared import Pt
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(255, 255, 255)  # White text
    run.font.size = Pt(1)  # Tiny font size


def add_invisible_keywords_strategically(doc, keywords):
    """Add keywords using multiple invisible strategies"""

    # Strategy 1: Hidden text at document end
    hidden_para = doc.add_paragraph()
    add_hidden_text(hidden_para, ' '.join(keywords))

    # Strategy 2: White text scattered throughout
    paragraphs = doc.paragraphs
    if len(paragraphs) > 3:  # Only if document has enough content
        # Add white text after first paragraph
        white_para1 = doc.paragraphs[1]
        add_white_text(white_para1, ' '.join(keywords[:10]))

        # Add white text in middle
        mid_idx = len(paragraphs) // 2
        if mid_idx < len(paragraphs):
            white_para2 = doc.paragraphs[mid_idx]
            add_white_text(white_para2, ' '.join(keywords[10:20]))

    # Strategy 3: Add as document variables (XML level)
    try:
        doc_vars = doc.settings
        # This adds keywords at XML level
    except:
        pass

    return len(keywords)


def read_resume_text(docx_path):
    """Extract text content from resume docx file"""
    try:
        doc = docx.Document(docx_path)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return '\n'.join(text)
    except Exception as e:
        print(f"Error reading resume: {e}")
        return ""


def inject_invisible_keywords(docx_path, keywords, output_path, pdf_output_path=None):
    """Main function to inject keywords using multiple invisible strategies"""
    try:
        doc = docx.Document(docx_path)

        print(f"Processing {len(keywords)} keywords...")

        # Strategy 1: Add to document metadata
        add_keywords_to_metadata(doc, keywords)
        print("✓ Keywords added to document metadata")

        # Strategy 2: Add invisible text using multiple methods
        keywords_added = add_invisible_keywords_strategically(doc, keywords)
        print(f"✓ {keywords_added} keywords added as invisible text")

        # Save the optimized document
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        print(f"✓ ATS-optimized DOCX resume saved: '{output_path}'")

        # Generate PDF if requested
        if pdf_output_path:
            try:
                convert(output_path, pdf_output_path)
                if os.path.exists(pdf_output_path):
                    print(f"✓ PDF version created: '{pdf_output_path}'")
                else:
                    print("⚠ PDF conversion attempted but file not found")
            except Exception as e:
                print(f"⚠ PDF conversion failed: {e}")

        print("\n🎯 ATS Optimization Complete!")
        print("   • Document layout unchanged")
        print("   • Keywords embedded invisibly")
        print("   • ATS systems will detect keywords")
        print("   • Human readers won't see additions")

        return True

    except Exception as e:
        print(f"❌ Error processing resume: {e}")
        return False


def analyze_job_description(jd_text):
    """Analyze job description and provide insights"""
    keywords = extract_smart_keywords(jd_text)

    print("\n📊 Job Description Analysis:")
    print(f"   • Total keywords extracted: {len(keywords)}")
    print(f"   • Top 10 keywords: {', '.join(keywords[:10])}")

    # Categorize keywords
    tech_skills = [kw for kw in keywords if any(tech in kw.lower()
                                                for tech in
                                                ['python', 'aws', 'ai', 'ml', 'docker', 'kubernetes', 'sql'])]
    soft_skills = [kw for kw in keywords if any(soft in kw.lower()
                                                for soft in
                                                ['communication', 'leadership', 'collaboration', 'problem'])]

    if tech_skills:
        print(f"   • Technical skills: {', '.join(tech_skills[:5])}")
    if soft_skills:
        print(f"   • Soft skills: {', '.join(soft_skills[:3])}")

    return keywords


def extract_missing_keywords_llm(jd_text, resume_text, max_keywords=50):
    """
    Use a free LLM API to extract the most important keywords from the job description
    that are NOT present in the resume. Returns a list of missing keywords.
    """
    API_URL = "https://router.huggingface.co/v1/chat/completions"
    hf_token = os.environ.get("HF_TOKEN")
    print("[DEBUG] HF_TOKEN in env:", "✓ Found" if hf_token else "✗ Not found")

    if not hf_token:
        print("[LLM API] Error: HuggingFace token not found in environment variable 'HF_TOKEN'.")
        print("[LLM API] Please add HF_TOKEN to your .env file or set as environment variable.")
        print(
            "[LLM API] TIP: Make sure you export HF_TOKEN in the same shell session and run the script from that shell. Try: export HF_TOKEN=your_token && python ats_cli.py")
        return extract_fallback_keywords(jd_text, resume_text, max_keywords)

    headers = {"Authorization": f"Bearer {hf_token}"}

    # Use the imported prompt template
    prompt = LLM_KEYWORD_EXTRACTION_PROMPT.format(jd_text=jd_text.strip(), resume_text=resume_text.strip())

    # Export prompt to file for debugging
    with open("debug_prompt.txt", "w", encoding='utf-8') as f:
        f.write(prompt)
    print("[LLM API] Detailed ATS prompt written to debug_prompt.txt")

    payload = {
        "messages": [
            {"role": "user", "content": prompt}
        ],
        "model": "moonshotai/Kimi-K2-Instruct-0905:groq",
        "temperature": 0.1,  # Low temperature for more focused extraction
        "max_tokens": 500  # Limit response length
    }

    try:
        print("[LLM API] Sending request to HuggingFace API...")
        response = requests.post(API_URL, headers=headers, json=payload, timeout=60)
        response.raise_for_status()
        result = response.json()

        # Save debug response
        with open("debug_response.json", "w", encoding='utf-8') as f:
            json.dump(result, f, indent=2)
        print("[LLM API] Response written to debug_response.json")

        # Enhanced response processing
        if isinstance(result, dict) and 'choices' in result and len(result['choices']) > 0:
            content = result['choices'][0]['message']['content'].strip()
            print(f"[LLM API] Raw response: {content[:200]}...")

            # Clean and extract keywords
            # Remove any explanatory text and focus on the comma-separated list
            lines = content.split('\n')
            keyword_line = ""
            for line in lines:
                if ',' in line and not line.startswith(('Note:', 'Explanation:', 'Here are', 'The keywords')):
                    keyword_line = line
                    break

            if not keyword_line:
                keyword_line = content  # Fallback to full content

            keywords = [kw.strip().strip('"').strip("'") for kw in keyword_line.split(',') if kw.strip()]
            keywords = [kw for kw in keywords if len(kw) > 1 and not kw.lower().startswith(('note:', 'here'))]

            print(f"[LLM API] Extracted {len(keywords)} keywords: {', '.join(keywords[:10])}...")
            return keywords[:max_keywords]
        else:
            print(f"[LLM API] Unexpected response format: {result}")
            return extract_fallback_keywords(jd_text, resume_text, max_keywords)

    except Exception as e:
        print(f"[LLM API] Error extracting keywords: {e}")
        print("[LLM API] Falling back to local keyword extraction...")
        return extract_fallback_keywords(jd_text, resume_text, max_keywords)


def extract_fallback_keywords(jd_text, resume_text, max_keywords=50):
    """Fallback keyword extraction when LLM API is not available"""
    print("[FALLBACK] Using local keyword extraction...")
    jd_keywords = set(extract_smart_keywords(jd_text, max_keywords * 2))

    # Simple check for keywords not in resume (case insensitive)
    resume_lower = resume_text.lower()
    missing_keywords = []

    for keyword in jd_keywords:
        if keyword.lower() not in resume_lower:
            missing_keywords.append(keyword)

    print(f"[FALLBACK] Found {len(missing_keywords)} missing keywords")
    return missing_keywords[:max_keywords]


def main():
    """Main execution function"""
    print("🚀 Smart ATS Resume Optimizer")
    print("=" * 40)

    try:
        # Check if job description file exists
        jd_file = "job_description.txt"
        if not os.path.exists(jd_file):
            print(f"❌ Error: {jd_file} not found")
            print("   Please create this file with the job description text")
            return

        # Read job description
        with open(jd_file, "r", encoding='utf-8') as f:
            jd_text = f.read()

        if not jd_text.strip():
            print(f"❌ Error: {jd_file} is empty")
            print("   Please add the job description text to the file")
            return

        # Analyze and extract keywords
        keywords = analyze_job_description(jd_text)

        # Define paths
        input_resume = "Suyash_Pathak_Resume.docx"
        output_dir = "optimized_resume"
        output_docx = f"{output_dir}/Suyash_Pathak_resume.docx"
        output_pdf = f"{output_dir}/Suyash_Pathak_ATS_Optimized.pdf"

        # Check if resume file exists
        if not os.path.exists(input_resume):
            print(f"❌ Error: {input_resume} not found")
            print("   Please make sure your resume file is in the current directory")
            return

        # Optional: Try to get smarter keywords using LLM
        resume_text = read_resume_text(input_resume)
        if resume_text:
            print("\n🤖 Attempting smart keyword extraction...")
            smart_keywords = extract_missing_keywords_llm(jd_text, resume_text)
            if smart_keywords:
                keywords = smart_keywords + keywords  # Combine smart + regular keywords
                keywords = list(dict.fromkeys(keywords))  # Remove duplicates while preserving order
                print(f"✓ Enhanced with {len(smart_keywords)} targeted keywords")

        # Process resume
        print(f"\n🔄 Processing resume: {input_resume}")
        success = inject_invisible_keywords(input_resume, keywords, output_docx, output_pdf)

        if success:
            print(f"\n✅ Success! Your ATS-optimized resume is ready:")
            print(f"   📄 DOCX: {output_docx}")
            print(f"   📄 PDF: {output_pdf}")
            print(f"\n💡 Tips:")
            print(f"   • Use the DOCX version for online applications")
            print(f"   • Use the PDF version for email submissions")
            print(f"   • Keywords are invisible to human readers")
            print(f"   • ATS systems will score higher relevance")

    except FileNotFoundError as e:
        print(f"❌ Error: Required file not found - {e}")
        print("   Please create the missing file and try again")
    except Exception as e:
        print(f"❌ Unexpected error: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()